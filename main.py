from fastapi import FastAPI, HTTPException, Response, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
from docx import Document
from io import BytesIO
import os
import datetime
import PyPDF2

# Initialize FastAPI app
app = FastAPI()

# Load API key
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY is not set. Please add it to your environment variables.")

client = Groq(api_key=GROQ_API_KEY)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class QuizRequest(BaseModel):
    text: str
    num_questions: int = 10

def extract_text_from_pdf(file: UploadFile) -> str:
    """Extract text from a PDF file with better error handling."""
    try:
        reader = PyPDF2.PdfReader(file.file)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            raise ValueError("No readable text found in the PDF.")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def preprocess_text(text: str) -> str:
    """Pre-process text for better question generation."""
    try:
        prompt = f"Normalize and simplify the following text for quiz question generation:\n{text}"
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": "You are an expert in simplifying text for educational quizzes."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Text preprocessing error: {str(e)}")

def select_sentences(text: str, num_questions: int) -> list:
    """Select key sentences to generate questions."""
    try:
        prompt = f"Select {num_questions} important sentences from this text to create quiz questions:\n{text}"
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": f"You are skilled at identifying key sentences for quiz generation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300
        )
        return response.choices[0].message.content.strip().split("\n")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sentence selection error: {str(e)}")

def generate_question(sentence: str, question_type: str) -> dict:
    """Generate structured quiz question from a sentence."""
    try:
        prompt = f"Create a {question_type} question from this sentence:\n{sentence}"
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": f"You are an expert in generating {question_type} quiz questions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200
        )
        return {"question": response.choices[0].message.content.strip(), "type": question_type}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Question generation error: {str(e)}")

def generate_quiz(text: str, num_questions: int) -> list:
    """Generate a full structured quiz with various question types."""
    try:
        preprocessed_text = preprocess_text(text)
        sentences = select_sentences(preprocessed_text, num_questions)
        quiz = []
        question_types = ["multiple-choice", "true/false", "question/answer", "fill in the blanks"]

        for sentence in sentences:
            for q_type in question_types:
                question = generate_question(sentence, q_type)
                if question.get("question"):  # Ensure valid question
                    quiz.append(question)
        return quiz
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Quiz generation failed: {str(e)}")

def create_quiz_document(quiz: list) -> BytesIO:
    """Generate a downloadable DOCX quiz document."""
    try:
        doc = Document()
        doc.add_heading('Generated Quiz', 0)
        
        for i, question in enumerate(quiz, 1):
            doc.add_paragraph(f"Q{i}: {question['question']}")
            doc.add_paragraph(f"Type: {question['type']}\n")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating document: {str(e)}")

@app.post("/generate/")
async def generate_quiz_endpoint(request: QuizRequest):
    """Generate and return a structured quiz."""
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        if request.num_questions < 1:
            raise HTTPException(status_code=400, detail="Number of questions must be at least 1")

        quiz = generate_quiz(request.text, request.num_questions)
        return {"quiz": quiz}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download/")
async def download_quiz_endpoint(request: QuizRequest):
    """Generate and return a downloadable DOCX quiz."""
    try:
        quiz = generate_quiz(request.text, request.num_questions)
        buffer = create_quiz_document(quiz)
        filename = f"quiz_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-pdf/")
async def upload_pdf(file: UploadFile = File(...)):
    """Extract text from PDF and generate a quiz document."""
    try:
        text = extract_text_from_pdf(file)
        quiz = generate_quiz(text, 10)  # Default to 10 questions
        buffer = create_quiz_document(quiz)
        filename = f"quiz_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def welcome():
    """Welcome message for API root endpoint."""
    return {"message": "Welcome to the Quiz Generation API!"}
